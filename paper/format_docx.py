"""
Format final_report.docx theo chuẩn học thuật Việt Nam:
- Font: Times New Roman 13pt cho body, Times New Roman cho heading (size khác nhau)
- Line spacing: 1.5
- Margins: 2.5cm trái, 2cm phải, 2cm trên, 2cm dưới
- Justify alignment cho body
- Heading 1, 2, 3 đậm
- Đánh số trang ở footer
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT = "/sessions/amazing-kind-allen/mnt/Master-Study-IDT/Pizza4Ps_DataMining/paper/final_report.docx"
OUTPUT = "/sessions/amazing-kind-allen/mnt/Master-Study-IDT/Pizza4Ps_DataMining/paper/final_report.docx"

doc = Document(INPUT)

# ---------------- Page setup ----------------
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

# ---------------- Default style ----------------
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(13)
# Set East Asia font (cần thiết khi gặp ký tự tiếng Việt unicode)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.find(qn("w:rFonts"))
if rfonts is None:
    rfonts = OxmlElement("w:rFonts")
    rpr.append(rfonts)
rfonts.set(qn("w:ascii"), "Times New Roman")
rfonts.set(qn("w:hAnsi"), "Times New Roman")
rfonts.set(qn("w:eastAsia"), "Times New Roman")
rfonts.set(qn("w:cs"), "Times New Roman")

style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(6)

# ---------------- Heading styles ----------------
heading_sizes = {
    "Heading 1": 16,
    "Heading 2": 14,
    "Heading 3": 13,
    "Heading 4": 13,
    "Title": 18,
}
for name, size in heading_sizes.items():
    try:
        s = doc.styles[name]
        s.font.name = "Times New Roman"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0, 0, 0)
        rpr = s.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "Times New Roman")
        s.paragraph_format.line_spacing = 1.5
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.keep_with_next = True
    except KeyError:
        pass

# ---------------- Format body paragraphs ----------------
for para in doc.paragraphs:
    if para.style.name == "Normal":
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.first_line_indent = Cm(0.5)
        para.paragraph_format.space_after = Pt(6)
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)
            r = run._element
            rpr = r.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:ascii"), "Times New Roman")
            rfonts.set(qn("w:hAnsi"), "Times New Roman")
            rfonts.set(qn("w:eastAsia"), "Times New Roman")
    elif para.style.name.startswith("Heading"):
        # Đảm bảo runs trong heading dùng Times New Roman
        for run in para.runs:
            run.font.name = "Times New Roman"
            r = run._element
            rpr = r.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:ascii"), "Times New Roman")
            rfonts.set(qn("w:hAnsi"), "Times New Roman")
            rfonts.set(qn("w:eastAsia"), "Times New Roman")

# ---------------- Format tables ----------------
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.line_spacing = 1.15
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)  # tables hơi nhỏ hơn body
                    r = run._element
                    rpr = r.get_or_add_rPr()
                    rfonts = rpr.find(qn("w:rFonts"))
                    if rfonts is None:
                        rfonts = OxmlElement("w:rFonts")
                        rpr.append(rfonts)
                    rfonts.set(qn("w:ascii"), "Times New Roman")
                    rfonts.set(qn("w:hAnsi"), "Times New Roman")
                    rfonts.set(qn("w:eastAsia"), "Times New Roman")

# ---------------- Page numbering in footer ----------------
def add_page_number(footer):
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    run = p.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

for section in doc.sections:
    add_page_number(section.footer)

doc.save(OUTPUT)
print(f"✓ Đã format xong: {OUTPUT}")

# Đếm thử số word
total_chars = 0
for para in doc.paragraphs:
    total_chars += len(para.text)
print(f"Tổng số ký tự body: {total_chars:,}")
print(f"Số đoạn văn: {len(doc.paragraphs)}")
print(f"Số bảng: {len(doc.tables)}")
